"""
File content extraction utility for PDF, DOCX, and image files
"""
import io
import os
from typing import Optional
import requests
from pathlib import Path

try:
    import PyPDF2
    PDF_AVAILABLE = True
except ImportError:
    PDF_AVAILABLE = False

try:
    from docx import Document
    DOCX_AVAILABLE = True
except ImportError:
    DOCX_AVAILABLE = False

try:
    from PIL import Image
    try:
        import pytesseract
        OCR_AVAILABLE = True
    except ImportError:
        OCR_AVAILABLE = False
except ImportError:
    OCR_AVAILABLE = False


def extract_text_from_pdf(file_content: bytes) -> str:
    """Extract text from PDF file"""
    if not PDF_AVAILABLE:
        return ""  # Return empty string if PDF extraction not available
    
    try:
        pdf_file = io.BytesIO(file_content)
        pdf_reader = PyPDF2.PdfReader(pdf_file)
        text = ""
        
        for page_num in range(len(pdf_reader.pages)):
            page = pdf_reader.pages[page_num]
            text += page.extract_text() + "\n"
        
        return text.strip() if text else ""
    except Exception as e:
        # Return empty string on error (file might be corrupted or image-only PDF)
        print(f"Error extracting PDF text: {str(e)}")
        return ""


def extract_text_from_docx(file_content: bytes) -> str:
    """Extract text from DOCX file"""
    if not DOCX_AVAILABLE:
        return ""  # Return empty string if DOCX extraction not available
    
    try:
        docx_file = io.BytesIO(file_content)
        doc = Document(docx_file)
        text = ""
        
        for paragraph in doc.paragraphs:
            text += paragraph.text + "\n"
        
        # Also extract text from tables
        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    text += cell.text + " "
                text += "\n"
        
        return text.strip() if text else ""
    except Exception as e:
        # Return empty string on error (file might be corrupted)
        print(f"Error extracting DOCX text: {str(e)}")
        return ""


def extract_text_from_image(file_content: bytes, file_extension: str) -> str:
    """Extract text from image using OCR"""
    if not OCR_AVAILABLE:
        return ""  # Return empty string if OCR not available
    
    try:
        import pytesseract
        image = Image.open(io.BytesIO(file_content))
        text = pytesseract.image_to_string(image)
        return text.strip() if text else ""
    except Exception as e:
        # OCR failed, return empty string (not an error, just no text found)
        print(f"Error extracting text from image: {str(e)}")
        return ""


def extract_text_from_file(file_content: bytes, filename: str, content_type: Optional[str] = None) -> str:
    """
    Extract text content from uploaded file
    
    Args:
        file_content: File content as bytes
        filename: Original filename
        content_type: MIME type of the file
    
    Returns:
        Extracted text content (empty string if extraction fails or not supported)
    """
    if not filename:
        return ""
    
    file_extension = Path(filename).suffix.lower()
    
    # Determine file type and extract accordingly
    if file_extension == '.pdf' or (content_type and 'pdf' in content_type.lower()):
        return extract_text_from_pdf(file_content)
    
    elif file_extension in ['.docx', '.doc'] or (content_type and 'word' in content_type.lower()):
        if file_extension == '.docx':
            return extract_text_from_docx(file_content)
        else:
            return ""  # DOC files not supported, return empty string
    
    elif file_extension in ['.jpg', '.jpeg', '.png', '.gif', '.bmp'] or (content_type and 'image' in content_type.lower()):
        return extract_text_from_image(file_content, file_extension)
    
    else:
        return ""  # Return empty string for unsupported file types


def extract_text_from_url(file_url: str) -> str:
    """
    Download file from URL and extract text
    
    Args:
        file_url: URL of the file
    
    Returns:
        Extracted text content (empty string if extraction fails)
    """
    try:
        response = requests.get(file_url, timeout=30, stream=True)
        response.raise_for_status()
        
        # Get file extension from URL
        filename = os.path.basename(file_url.split('?')[0])  # Remove query parameters
        content_type = response.headers.get('content-type', '')
        
        # Read file content
        file_content = response.content
        
        return extract_text_from_file(file_content, filename, content_type)
    except Exception as e:
        print(f"Error downloading and extracting from URL: {str(e)}")
        return ""  # Return empty string on error
